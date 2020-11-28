import io
from collections import defaultdict
from tempfile import NamedTemporaryFile
import matplotlib.pyplot as plt
import numpy as np
import textract
import cv2
import pytesseract

from speech2text import speech_to_text
from text_process import change_text_person
from docx import Document
from docx.shared import Inches
import streamlit as st
from streamlit.uploaded_file_manager import UploadedFile

default_return = {
    'already_processed': False,

}


# @st.cache(hash_funcs={UploadedFile: lambda x: x.name}, show_spinner=False)
def read_file(datastream):
    with st.spinner('Обработка данных...'):
        with NamedTemporaryFile(delete=False, suffix=datastream.name) as f:
            datastream.seek(0)
            f.write(datastream.read())
            datastream.name = datastream.name.lower()
            if datastream.name.endswith('.jpeg') or datastream.name.endswith('.jpg') or datastream.name.endswith(
                    '.png'):
                return read_image(f.name)
            elif datastream.name.endswith('.docx'):
                return read_docx(f.name)
            elif datastream.name.endswith('.m4a') or datastream.name.endswith('.wav') \
                    or datastream.name.endswith('.ogg') or datastream.name.endswith('.mp3'):
                return_dict = default_return.copy()
                return_dict['audio'] = f.name
                return_dict['text'] = speech_to_text(f.name)
                return return_dict
            else:
                return read_text(f.name)


def read_text(file_name):
    return_dict = default_return.copy()
    text = textract.process(file_name).decode("utf8")
    return_dict['text'] = text
    return return_dict


def read_image(file_name):
    return_dict = default_return.copy()
    gray = cv2.imread(file_name, 0)
    sharpen = cv2.GaussianBlur(gray, (0, 0), 3)
    sharpen = cv2.addWeighted(gray, 1.5, sharpen, -0.5, 0)

    # apply adaptive threshold to get black and white effect
    th3 = cv2.adaptiveThreshold(sharpen, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 21, 15)
    return_dict['image'] = th3
    return_dict['text'] = pytesseract.image_to_string(th3, lang='rus')
    return return_dict


def read_docx(file_name):
    return_dict = default_return.copy()
    document = Document(file_name)
    tokens = []
    ners = defaultdict(list)
    for paragraph in document.paragraphs:
        new_tokens, new_ners = change_text_person(paragraph.text)
        for key, value in new_ners.items():
            ners[key] += value
        count_changes = np.sum([isinstance(token, tuple) for token in new_tokens])
        if len(tokens) > 0:
            tokens.append('\n')
        tokens.extend(new_tokens)
        paragraph.text = ' '.join([token[0] if isinstance(token, tuple) else token for token in new_tokens])

    return_dict['already_processed'] = True
    return_dict['text'] = tokens
    return_dict['ners'] = ners
    target_stream = io.BytesIO()
    document.save(target_stream)
    return_dict['saved_doc'] = target_stream
    return return_dict


def write_docx(text, template=None):
    paragraph = None
    if template is not None:
        with NamedTemporaryFile(delete=True, suffix=template.name) as f:
            template.seek(0)
            f.write(template.read())
            document = Document(f.name)
        for p in document.paragraphs:
            if 'МЕСТО ДЛЯ ТЕКСТА' in p.text.upper():
                paragraph = p
        if paragraph is None:
            paragraph = document.add_paragraph()
    else:
        document = Document()
        paragraph = document.add_paragraph()

    paragraph.text = text

    target_stream = io.BytesIO()
    document.save(target_stream)
    return target_stream
